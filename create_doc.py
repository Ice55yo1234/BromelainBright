import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

doc = docx.Document()

# Page Margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Color Palette (Forest Green & Sage theme)
COLOR_PRIMARY = RGBColor(27, 67, 50)     # #1B4332 Forest Green
COLOR_SECONDARY = RGBColor(82, 183, 136) # #52B788 Sage Green
COLOR_TEXT = RGBColor(43, 45, 66)        # Dark Charcoal
COLOR_MUTED = RGBColor(108, 117, 125)    # Gray

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_heading_1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Cordia New'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    return p

def add_heading_2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Cordia New'
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = COLOR_SECONDARY
    return p

def add_body_p(text, bold_prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Cordia New'
        r_pre.font.size = Pt(14)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_TEXT
    r = p.add_run(text)
    r.font.name = 'Cordia New'
    r.font.size = Pt(14)
    r.font.color.rgb = COLOR_TEXT
    return p

def add_bullet(bold_prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix + " ")
        r_pre.font.name = 'Cordia New'
        r_pre.font.size = Pt(14)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_PRIMARY
    r = p.add_run(text)
    r.font.name = 'Cordia New'
    r.font.size = Pt(14)
    r.font.color.rgb = COLOR_TEXT
    return p

# --- TITLE & OVERVIEW ---
title_p = doc.add_paragraph()
title_p.paragraph_format.space_after = Pt(2)
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
t_run = title_p.add_run("เอกสารสรุปรายละเอียดฟีเจอร์และแนวทางการพัฒนาเว็บ")
t_run.font.name = 'Cordia New'
t_run.font.size = Pt(22)
t_run.font.bold = True
t_run.font.color.rgb = COLOR_PRIMARY

sub_p = doc.add_paragraph()
sub_p.paragraph_format.space_after = Pt(16)
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
s_run = sub_p.add_run("Bromelain Bright — Luxury Sustainable Skincare Startup")
s_run.font.name = 'Cordia New'
s_run.font.size = Pt(16)
s_run.font.bold = True
s_run.font.color.rgb = COLOR_SECONDARY

# Overview Section
add_heading_1("1. ภาพรวมของเว็บไซต์ (Executive Summary)")
add_body_p("เว็บไซต์ Bromelain Bright เป็นเว็บสตาร์ทอัพระดับพรีเมียมสำหรับแบรนด์สกินแคร์ความงามยั่งยืน (Upcycled Skincare) ที่สกัดเอนไซม์บรอมมีเลน (Bromelain Enzyme) จากแกนสับปะรดเหลือทิ้ง โดยได้รับการพัฒนาขึ้นมาครบถ้วนทั้งด้าน UI/UX ดีไซน์หรูหรา และระบบหลังบ้าน (Backend Demo API)")
add_bullet("วัตถุประสงค์หลัก:", "ใช้เป็นเว็บนำเสนอในการประกวดสตาร์ทอัพ (เช่น ISKKU x HEX) และใช้เป็นแพลตฟอร์มอีคอมเมิร์ซเต็มรูปแบบ")
add_bullet("เทคโนโลยีที่ใช้:", "HTML5, CSS3 (Custom Design Tokens), Vanilla JS (ES6), Node.js, Express.js API, SQLite Database")
add_bullet("ข้อมูลติดต่อแบรนด์:", "rapeephat.ph@kkumail.com | โทร +66 94-638-8671 | มหาวิทยาลัยขอนแก่น วิทยาเขตหนองคาย (Khon Kaen University, Nong Khai Campus)")

# Summary Table of All 19 Pages
add_heading_1("2. สรุปรายละเอียดหน้าที่ความสามารถของทุกหน้า (All 19 Pages)")

pages_data = [
    ("1. หน้าแรก (index.html)", "Hero Banner, แสดงผลิตภัณฑ์เด่น, นวัตกรรมเอนไซม์, ตัวเลขอิมแพกต์, รีวิวผู้ใช้, Gallery, AI Preview", "เชื่อมต่อ AI Real-time, ระบบแนะนำสินค้าตามพฤติกรรมผู้ใช้ (Personalization)"),
    ("2. คอลเลกชันสินค้า (products.html)", "สรุปแนวคิด Product Line (Sachet, Jar, Essence) พร้อมตารางเปรียบเทียบคุณสมบัติสินค้าแบบละเอียด", "ระบบฟิลเตอร์เปรียบเทียบเทียบกับแบรนด์คู่แข่งในตลาด"),
    ("3. สินค้าแบบซอง (product-sachet.html)", "หน้ารายละเอียด Trial Sachet (25ml x 5 ซอง / ฿290) แสดงวิธีใช้ ส่วนผสม และปุ่มกดสั่งซื้อ/Wishlist", "เลือกจำนวนเซ็ตแบบ Custom, ระบบแจ้งเตือนสินค้าพร้อมส่ง"),
    ("4. สินค้าแบบกระปุก (product-jar.html)", "หน้ารายละเอียด Luxury Jar (100g / ฿890) เน้นบรรจุภัณฑ์แก้วรีไซเคิลและความคุ้มค่า", "ระบบสมัครสมาชิกรายเดือนเพื่อส่งสินค้าอัตโนมัติ (Subscription System)"),
    ("5. สินค้าเอสเซนส์ (product-essence.html)", "หน้ารายละเอียด Concentrated Essence (30ml / ฿1,290) เซรั่มเข้มข้นพร้อม Dropper", "ระบบ AR แสดงการหยดเอสเซนส์ลงผิวจริงผ่านกล้อง"),
    ("6. ร้านค้าออนไลน์ (shop.html)", "รายการสินค้าทั้งหมด มีระบบค้นหา (Search Real-time), กรองหมวดหมู่ (Category Filter) และจัดเรียงราคา", "รองรับระบบกรองตามปัญหาผิว (Acne, Dark spots, Aging)"),
    ("7. ตะกร้าสินค้า (cart.html)", "คำนวณราคาสินค้า ปรับจำนวน ชำระค่าจัดส่งฟรีเมื่อครบกำหนด และคำนวณราคารวมอัตโนมัติ", "ระบบคูปองส่วนลด (Promo Codes) และการแถมของขวัญตามยอดสั่งซื้อ"),
    ("8. ชำระเงิน (checkout.html)", "ฟอร์มที่อยู่ละเอียด (ไทย/ต่างประเทศ) พร้อมระบบชำระเงินครบ (PromptPay, บัตรเครดิต, PayPal, COD, Mobile Bank)", "เชื่อมต่อ Payment Gateway จริง (Omise, Stripe) และระบบติดตามพัสดุ (Kerry/DHL API)"),
    ("9. งานวิจัย (research.html)", "ข้อมูลทางวิทยาศาสตร์การสกัดเอนไซม์ Bromelain กราฟแสดงผลลัพธ์จาก Clinical Trial และ FAQs", "ดาวน์โหลดเอกสารวิจัย Whitepaper (PDF), ระบบสอบถามผู้เชี่ยวชาญ"),
    ("10. ความยั่งยืน (sustainability.html)", "Dashboard แสดงตัวเลขอิมแพกต์ลดขยะสับปะรด ลด CO₂ และโมเดลเศรษฐกิจ BCG (Bio-Circular-Green)", "คำนวณ Carbon Footprint Offset Real-time ที่ผู้ใช้ช่วยลดได้"),
    ("11. ตรวจสอบที่มา (traceability.html)", "Interactive Timeline แสดงเส้นทาง Farm-to-Face จากฟาร์มสับปะรดไทยจนถึงมือผู้บริโภค", "ระบบสแกน QR Code บนกล่องสินค้าเพื่อดูฟาร์มต้นกำเนิดจริง (Blockchain)"),
    ("12. ระบบ AI สภาพผิว (ai-analysis.html)", "จำลองการอัปโหลดรูปถ่ายเพื่อวิเคราะห์ปัญหาผิว และประมวลผลแนะนำสินค้าที่เหมาะสม", "เชื่อมต่อ AI Vision Model จริง (เช่น TensorFlow/OpenAI) วิเคราะห์ริ้วรอยและความชุ่มชื้น"),
    ("13. เกี่ยวกับบริษัท (company.html)", "เรื่องราวผู้ก่อตั้งจาก มข. วิทยาเขตหนองคาย, วิสัยทัศน์, รางวัล ISKKU x HEX และ Timeline การเติบโต", "หน้าข่าวสารและงานแถลงข่าวของบริษัท (Press Release / Media Kit)"),
    ("14. ความเป็นมาแบรนด์ (about.html)", "เจตนารมณ์ ปรัชญาความงามยั่งยืน และพันธกิจการช่วยเหลือเกษตรกรไทย", "วิดีโอสัมภาษณ์เกษตรกรผู้ปลูกสับปะรดและความเป็นอยู่ที่ดีขึ้น"),
    ("15. ติดต่อเรา (contact.html)", "แบบฟอร์มติดต่อ อีเมล (rapeephat.ph@kkumail.com), เบอร์โทรศัพท์, โซเชียลมีเดีย และ Google Maps แผนที่จริง", "ระบบ Live Chat สนทนาสดกับฝ่ายบริการลูกค้า"),
    ("16. เข้าสู่ระบบ (login.html)", "ฟอร์ม Sign in ด้วย Glassmorphism UI พร้อมปุ่ม Auto-fill Demo Account สำหรับการนำเสนอ", "รองรับ Social Login (Google, Facebook, Apple ID, LINE Login)"),
    ("17. สมัครสมาชิก (register.html)", "ฟอร์มสร้างบัญชีผู้ใช้ใหม่ พร้อมตรวจสอบความถูกต้องของรหัสผ่านและความปลอดภัย", "ระบบยืนยันตัวตนผ่าน OTP ทาง SMS/Email"),
    ("18. แผงควบคุมผู้ใช้ (dashboard.html)", "จัดการข้อมูลส่วนตัว ดูรายการ Wishlist สินค้าที่ชอบ และประวัติการสั่งซื้อ", "ระบบสะสมแต้ม Reward Points, ระดับสมาชิก (Tiering & Loyalty Club)"),
    ("19. หน้าไม่พบข้อมูล (404.html)", "หน้า 404 ดีไซน์สวยงามเข้ากับแบรนด์ พร้อมปุ่มนำทางกลับหน้าหลัก", "ระบบแนะนำหน้ายอดนิยมเพื่อดึงผู้ใช้กลับเข้าสู่เว็บไซต์")
]

table = doc.add_table(rows=1, cols=3)
table.autofit = False

# Header Row
hdr_cells = table.rows[0].cells
headers = ['ชื่อหน้าเว็บ', 'ฟังก์ชันการทำงานปัจจุบัน (Current Capabilities)', 'แนวทางการพัฒนาต่อยอด (Future Roadmap)']
for i, title_text in enumerate(headers):
    hdr_cells[i].text = title_text
    set_cell_background(hdr_cells[i], "1B4332")
    p = hdr_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = 'Cordia New'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

# Set Column Widths
col_widths = [Inches(1.8), Inches(2.7), Inches(2.2)]

# Populate Data Rows
for row_idx, data in enumerate(pages_data):
    row_cells = table.add_row().cells
    bg_color = "F9FBF9" if row_idx % 2 == 1 else "FFFFFF"
    for i in range(3):
        row_cells[i].text = data[i]
        set_cell_background(row_cells[i], bg_color)
        p = row_cells[i].paragraphs[0]
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.name = 'Cordia New'
            run.font.size = Pt(12)
            run.font.color.rgb = COLOR_TEXT

# Apply width to all table cells
for row in table.rows:
    for i, w in enumerate(col_widths):
        row.cells[i].width = w

# --- FUTURE EXPANSION SUMMARY ---
add_heading_1("3. สรุปศักยภาพการนำไปต่อยอดในอนาคต (Future Technical Roadmap)")
add_bullet("การเชื่อมต่อระบบชำระเงินจริง (Payment Gateway):", "สามารถต่อยอดเชื่อม Stripe, Omise, หรือ GB Prime Pay เพื่อรับเงินจริงผ่าน PromptPay และบัตรเครดิตทั่วโลก")
add_bullet("การนำ AI มาใช้จริง (Production AI Skin Analysis):", "นำโมเดล Computer Vision วิเคราะห์ภาพใบหน้าจริง เพื่อประมวลผลความชุ่มชื้น สิว และริ้วรอย พร้อมจัดเซ็ตสินค้าเฉพาะบุคคล")
add_bullet("ระบบ Blockchain Traceability:", "เชื่อมต่อข้อมูลกับ Smart Contract เพื่อให้ลูกค้าสแกน QR Code บนกล่องสินค้า ตรวจสอบย้อนกลับไปถึงเกษตรกรผู้ปลูกสับปะรดในไทย")
add_bullet("ระบบฐานข้อมูลและการจัดการหลังบ้าน (Full Backend & CMS):", "พัฒนาต่อยอดจาก Express Demo เป็นระบบหลังบ้านเต็มรูปแบบเพื่อจัดการสต็อก การจัดส่งพัสดุอัตโนมัติ และบัญชีผู้ใช้")

# Save document
output_path = r"c:\antigravity\ข้าวโพดอบกรอบ\Bromelain_Bright_Website_Summary_v2.docx"
doc.save(output_path)
print("SUCCESS: Document created.")
